from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
import openai, os
from docx import Document
from io import BytesIO
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv

load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

router = APIRouter()

class ResumeRequest(BaseModel):
    name: str
    job_description: str

@router.post("/generate-resume")
def generate_resume(data: ResumeRequest):
    try:
        prompt = f"""
        Create a professional resume for {data.name}, tailored to this job description:
        {data.job_description}
        Include Summary, Skills, Experience, and Education sections.
        """

        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}]
        )

        content = response.choices[0].message.content

        # Generate .docx
        doc = Document()
        doc.add_heading(f"{data.name} - AI Resume", 0)
        for line in content.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={data.name.replace(' ', '_')}_resume.docx"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

from fastapi import APIRouter
from fastapi.responses import FileResponse
from utils.system.temp_storage_manager import load_temp_file
import os

router = APIRouter()

@router.get("/download/{filename}")
def download_resume(filename: str):
    file_path = f"/tmp/career_ai_vault/{filename}"
    if os.path.exists(file_path):
        return FileResponse(path=file_path, filename=filename, media_type="application/pdf")
    return {"error": "File not found."}
